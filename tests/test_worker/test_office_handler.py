from pathlib import Path

import pytest
from docx import Document

from worker import office_handler


@pytest.mark.asyncio
async def test_process_office_document_converts_docx(tmp_path: Path) -> None:
    document_path = tmp_path / "note.docx"
    document = Document()
    document.add_heading("Рабочая заметка", level=1)
    document.add_paragraph("Важный текст документа.")
    document.save(document_path)

    markdown = await office_handler.process_office_document(str(document_path))

    assert markdown is not None
    assert "Рабочая заметка" in markdown
    assert "Важный текст документа" in markdown


@pytest.mark.asyncio
async def test_process_office_document_returns_none_for_unsupported_file(tmp_path: Path) -> None:
    unsupported = tmp_path / "picture.jpg"
    unsupported.write_bytes(b"not an office document")

    assert await office_handler.process_office_document(str(unsupported)) is None


@pytest.mark.asyncio
async def test_process_office_document_returns_none_for_encrypted_document(monkeypatch: pytest.MonkeyPatch) -> None:
    def raise_encrypted(_path: str) -> str:
        raise office_handler.anydoc.EncryptedError("password required")

    monkeypatch.setattr(office_handler.anydoc, "to_markdown", raise_encrypted)

    assert await office_handler.process_office_document("protected.docx") is None
